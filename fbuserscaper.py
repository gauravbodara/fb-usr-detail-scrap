# -*- coding: utf-8 -*-
"""
@author: Gaurav Bodara
"""


#Import all these modules
from docx import Document
from selenium import webdriver
from selenium.webdriver.support.ui import WebDriverWait



#Initialize the document for writing to it the user data in word format
document = Document()


#Add heading
document.add_heading('User Details', 0)


#Open chrome and go to specified url 
url="https://www.facebook.com/"
driver=webdriver.Chrome()
driver.get(url)
driver.maximize_window()


#Facebook username and password
#Replace these with your own valid username and password
facebookUsername='username'
facebookPassword='password'


#For finding the username textfield and password textfield using xpath
emailFieldID='email'
passFieldID='pass'
loginButtonpath='//input[@value="Log In"]'
fbLogoXpath='(//a[contains(@href,"logo")])[1]'

#Find elements for login purpose and sends username and password in the textfield and find login button and clicks on it
emailFieldElem=WebDriverWait(driver,10).until(lambda driver: driver.find_element_by_id(emailFieldID))
passFieldElem=WebDriverWait(driver,10).until(lambda driver: driver.find_element_by_id(passFieldID))
loginButtonElem=WebDriverWait(driver,10).until(lambda driver: driver.find_element_by_xpath(loginButtonpath))
emailFieldElem.clear()
emailFieldElem.send_keys(facebookUsername)
passFieldElem.clear()
passFieldElem.send_keys(facebookPassword)
loginButtonElem.click()


#Replace name with the name you want to find user details 
user='name'
search_name = driver.find_element_by_name('q')
search_name.send_keys(user)
search_name.submit()
document.add_heading(user, level=1)


#Find Only one link that is first of all form the search results
link=WebDriverWait(driver,20).until(lambda driver: driver.find_element_by_xpath('//*[@id="xt_uniq_3"]/a'))
click_link=link.get_attribute("href").replace("?ref=br_rs","")
driver.get(click_link)


#for creating the link for user's about page and click on it
part_url=click_link
id=part_url.find("id")
if id==-1:
   about=driver.current_url+"/about"
else:
    part_url=url+part_url[id+3:]
    part_url=part_url.replace("&ref=br_rs","")
    about=part_url+"/about";
driver.get(about)


#for finding Education details of user and add it to word document
education=about+"/?section=education";
driver.get(education)
list_edu=WebDriverWait(driver,15).until(lambda driver: driver.find_element_by_id("pagelet_eduwork"));

document.add_heading("Education", level=1)
document.add_paragraph(list_edu.text)
document.add_paragraph("--------------------------------")

#for finding Living details of user and add it to word document
living=about+"/?section=living";
driver.get(living)
pagelet_hometown=WebDriverWait(driver,15).until(lambda driver: driver.find_element_by_id("pagelet_hometown"));

document.add_heading("Living", level=1)
document.add_paragraph(pagelet_hometown.text)
document.add_paragraph("--------------------------------")


#for finding Contact-details of user and add it to word document
contact=about+"/?section=contact-info"
driver.get(contact)
contact_info=WebDriverWait(driver,15).until(lambda driver: driver.find_element_by_id("pagelet_contact"));

document.add_heading("Contact_info", level=1)
document.add_paragraph(contact_info.text)
document.add_paragraph("--------------------------------")

#for finding Basic details of user and add it to word document
basic_contact=WebDriverWait(driver,15).until(lambda driver: driver.find_element_by_id("pagelet_basic"));

document.add_heading("Basic_contact", level=1)
document.add_paragraph(basic_contact.text)
document.add_paragraph("--------------------------------")


#for finding Relationship details of user and add it to word document
realtionships=about+"/?section=relationship"
driver.get(realtionships)
relations=WebDriverWait(driver,15).until(lambda driver: driver.find_element_by_id("pagelet_relationships"));

document.add_heading("Relations", level=1)
document.add_paragraph(relations.text)
document.add_paragraph("--------------------------------")


#for finding Bio details of user and add it to word document
details=about+"/?section=bio"
driver.get(details)
add_details=WebDriverWait(driver,15).until(lambda driver: driver.find_element_by_id("pagelet_bio"));

document.add_heading("Bio", level=1)
document.add_paragraph(add_details.text)
document.add_paragraph("--------------------------------")


#for finding Nickname details of user and add it to word document
nickname=WebDriverWait(driver,15).until(lambda driver: driver.find_element_by_id("pagelet_nicknames"));

document.add_heading("Nickname", level=1)
document.add_paragraph(nickname.text)
document.add_paragraph("--------------------------------")


##for finding quote of user and add it to word document
quote=WebDriverWait(driver,15).until(lambda driver: driver.find_element_by_id("pagelet_quotes"));

document.add_heading("Quote", level=1)
document.add_paragraph(quote.text)
document.add_paragraph("--------------------------------")

#for finding Life Events of user and add it to word document
life_event=about+"/?section=year-overviews"
driver.get(life_event)
life_events=WebDriverWait(driver,15).until(lambda driver: driver.find_element_by_class_name("_4qm1"));

document.add_heading("Life_events", level=1)
document.add_paragraph(life_events.text)
document.add_paragraph("--------------------------------")


#For saving document and name of it
filename=user+".docx";
document.save(filename)
print("-------------Successful-------------")


driver.quit()
